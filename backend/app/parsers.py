"""Document parsing. Returns extracted text and (for spreadsheets) lightweight table summaries.

Pure-Python deps only — no native compile chain required.
"""
from __future__ import annotations

import csv
import io
from pathlib import Path

import pdfplumber
from docx import Document
from openpyxl import load_workbook


MAX_TEXT_CHARS = 200_000  # hard cap per file to keep prompts manageable


def parse(path: Path, kind: str) -> tuple[str, list[dict]]:
    if kind == "pdf":
        return _pdf(path), []
    if kind == "docx":
        return _docx(path), []
    if kind == "xlsx":
        return _xlsx(path)
    raise ValueError(f"Unsupported file kind: {kind}")


def _pdf(path: Path) -> str:
    out = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            out.append(f"--- Page {i} ---\n{text}")
    return _trim("\n\n".join(out))


def _docx(path: Path) -> str:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for t in doc.tables:
        rows = []
        for row in t.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        tables.append("\n".join(rows))
    return _trim("\n\n".join(paragraphs + tables))


def _xlsx(path: Path) -> tuple[str, list[dict]]:
    wb = load_workbook(filename=str(path), data_only=True, read_only=True)
    text_parts: list[str] = []
    table_summaries: list[dict] = []
    try:
        for name in wb.sheetnames:
            ws = wb[name]
            rows = [
                r for r in ws.iter_rows(values_only=True)
                if any(c is not None and str(c).strip() != "" for c in r)
            ]
            if not rows:
                continue
            keep = _non_empty_columns(rows)
            filtered = [[r[i] if i < len(r) else None for i in keep] for r in rows]
            header = filtered[0]
            columns = [
                str(c) if c is not None else f"col_{i + 1}"
                for i, c in enumerate(header)
            ]
            n_rows = len(filtered) - 1  # exclude header row
            n_cols = len(columns)

            buf = io.StringIO()
            writer = csv.writer(buf)
            for r in filtered[:50]:
                writer.writerow(["" if c is None else c for c in r])
            preview = buf.getvalue()

            text_parts.append(
                f"--- Sheet: {name} ---\n"
                f"Shape: ({n_rows}, {n_cols})\n"
                f"Columns: {columns}\n\n{preview}"
            )
            table_summaries.append(
                {"sheet": name, "rows": n_rows, "cols": n_cols, "columns": columns}
            )
    finally:
        wb.close()
    return _trim("\n\n".join(text_parts)), table_summaries


def _non_empty_columns(rows: list[tuple]) -> list[int]:
    max_cols = max(len(r) for r in rows)
    has_value = [False] * max_cols
    for r in rows:
        for i in range(len(r)):
            if r[i] is not None and str(r[i]).strip() != "":
                has_value[i] = True
    return [i for i, v in enumerate(has_value) if v]


def _trim(s: str) -> str:
    return s if len(s) <= MAX_TEXT_CHARS else s[:MAX_TEXT_CHARS] + "\n... [truncated]"
